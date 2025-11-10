"""Text extraction service for parsing documents."""

import logging
from typing import Optional, Dict, Any
from pathlib import Path
import re

from sqlalchemy.ext.asyncio import AsyncSession

from app.domain.entities.preferences.example_resume import ExampleResume
from app.infrastructure.repositories.example_resume_repository import ExampleResumeRepository
from app.core.exceptions import TextExtractionError, ValidationException

logger = logging.getLogger(__name__)


class TextExtractionService:
    """Service for extracting text content from uploaded documents."""
    
    def __init__(self, db: AsyncSession):
        """
        Initialize text extraction service.
        
        Args:
            db: Database session
        """
        self.db = db
        self.repository = ExampleResumeRepository(db)

    async def extract_text_from_resume(self, resume_id: str) -> str:
        """
        Extract text content from an example resume.
        
        Args:
            resume_id: ExampleResume ID
            
        Returns:
            Extracted text content
            
        Raises:
            TextExtractionError: Extraction failed
            ValidationException: Resume not found or invalid
        """
        # Get resume from database
        resume = await self.repository.get_by_id(resume_id)
        if not resume:
            raise ValidationException(
                error_code="resume_not_found",
                message=f"Example resume not found: {resume_id}"
            )
        
        # Check if already extracted
        if resume.extracted_text and resume.text_extraction_status == "completed":
            return resume.extracted_text
        
        try:
            # Extract text based on file type
            file_extension = resume.file_metadata.file_extension.lower()
            storage_path = resume.storage_path
            
            if file_extension == "txt":
                extracted_text = await self._extract_from_txt(storage_path)
            elif file_extension == "pdf":
                extracted_text = await self._extract_from_pdf(storage_path)
            elif file_extension in ["doc", "docx"]:
                extracted_text = await self._extract_from_docx(storage_path)
            else:
                raise TextExtractionError(f"Unsupported file type: {file_extension}")
            
            # Clean and validate extracted text
            cleaned_text = self._clean_text(extracted_text)
            
            if len(cleaned_text.strip()) < 100:
                raise TextExtractionError("Extracted text is too short (< 100 characters)")
            
            # Update resume with extracted text
            await self.repository.update_text_extraction(
                resume_id=resume_id,
                extracted_text=cleaned_text,
                status="completed"
            )
            
            logger.info(f"Text extracted successfully: resume_id={resume_id}, length={len(cleaned_text)}")
            return cleaned_text
            
        except Exception as e:
            # Update status to failed
            await self.repository.update_text_extraction(
                resume_id=resume_id,
                extracted_text=None,
                status="failed",
                error=str(e)
            )
            
            logger.error(f"Text extraction failed: resume_id={resume_id}, error={e}")
            raise TextExtractionError(f"Failed to extract text: {str(e)}")

    async def extract_text_from_cover_letter(self, file_path: str, file_type: str) -> str:
        """
        Extract text content from a cover letter file.
        
        Args:
            file_path: Path to cover letter file
            file_type: File type/extension
            
        Returns:
            Extracted text content
            
        Raises:
            TextExtractionError: Extraction failed
        """
        try:
            file_type = file_type.lower()
            
            if file_type == "txt":
                extracted_text = await self._extract_from_txt(file_path)
            elif file_type == "pdf":
                extracted_text = await self._extract_from_pdf(file_path)
            elif file_type in ["doc", "docx"]:
                extracted_text = await self._extract_from_docx(file_path)
            else:
                raise TextExtractionError(f"Unsupported file type: {file_type}")
            
            # Clean and validate extracted text
            cleaned_text = self._clean_text(extracted_text)
            
            if len(cleaned_text.strip()) < 50:
                raise TextExtractionError("Cover letter text is too short (< 50 characters)")
            
            logger.info(f"Cover letter text extracted: path={file_path}, length={len(cleaned_text)}")
            return cleaned_text
            
        except Exception as e:
            logger.error(f"Cover letter text extraction failed: path={file_path}, error={e}")
            raise TextExtractionError(f"Failed to extract text: {str(e)}")

    async def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()

    async def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file.
        
        Note: This is a simplified implementation using PyPDF2.
        For production, consider using more robust libraries like pdfplumber or pdfminer.
        """
        try:
            import PyPDF2
            
            extracted_text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    extracted_text += page.extract_text() + "\n"
            
            return extracted_text
            
        except ImportError:
            # Fallback: Try to use basic text extraction
            logger.warning("PyPDF2 not installed, using basic extraction")
            return await self._extract_pdf_fallback(file_path)
        except Exception as e:
            raise TextExtractionError(f"PDF extraction failed: {str(e)}")

    async def _extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Note: This is a simplified implementation using python-docx.
        For production, consider handling more complex document structures.
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            extracted_text = ""
            
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + " "
                    extracted_text += "\n"
            
            return extracted_text
            
        except ImportError:
            logger.warning("python-docx not installed, cannot extract DOCX content")
            raise TextExtractionError("DOCX extraction not supported (python-docx not installed)")
        except Exception as e:
            raise TextExtractionError(f"DOCX extraction failed: {str(e)}")

    async def _extract_pdf_fallback(self, file_path: str) -> str:
        """Fallback PDF extraction method."""
        # This is a very basic fallback - in production you'd want a more robust solution
        logger.warning("Using fallback PDF extraction - results may be limited")
        return f"[PDF content from {Path(file_path).name}]\n\n" \
               "Note: Full PDF text extraction requires additional dependencies. " \
               "Please install PyPDF2 or pdfplumber for better PDF support."

    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple newlines to double newline
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
        text = re.sub(r'\n[ \t]+', '\n', text)  # Space at beginning of line
        
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)
        
        # Normalize common unicode characters
        text = text.replace('\u2019', "'")  # Right single quote
        text = text.replace('\u2018', "'")  # Left single quote
        text = text.replace('\u201c', '"')  # Left double quote
        text = text.replace('\u201d', '"')  # Right double quote
        text = text.replace('\u2013', '-')  # En dash
        text = text.replace('\u2014', '-')  # Em dash
        text = text.replace('\u2026', '...')  # Ellipsis
        
        # Remove extra whitespace at start/end
        text = text.strip()
        
        return text

    def get_extraction_stats(self, text: str) -> Dict[str, Any]:
        """
        Get statistics about extracted text.
        
        Args:
            text: Extracted text
            
        Returns:
            Dictionary with text statistics
        """
        if not text:
            return {"error": "No text provided"}
        
        # Basic stats
        char_count = len(text)
        word_count = len(text.split())
        line_count = text.count('\n') + 1
        
        # Estimate reading time (average 200 words per minute)
        reading_time_minutes = max(1, word_count // 200)
        
        # Check for common resume sections
        sections_found = []
        section_patterns = {
            "contact": r'\b(phone|email|address|linkedin|github)\b',
            "summary": r'\b(summary|objective|profile)\b',
            "experience": r'\b(experience|employment|work|position)\b',
            "education": r'\b(education|degree|university|college)\b',
            "skills": r'\b(skills|technologies|technical|tools)\b',
            "projects": r'\b(projects|portfolio|achievements)\b'
        }
        
        text_lower = text.lower()
        for section, pattern in section_patterns.items():
            if re.search(pattern, text_lower):
                sections_found.append(section)
        
        return {
            "character_count": char_count,
            "word_count": word_count,
            "line_count": line_count,
            "estimated_reading_time_minutes": reading_time_minutes,
            "sections_detected": sections_found,
            "extraction_quality": "good" if word_count > 100 and len(sections_found) > 2 else "fair" if word_count > 50 else "poor"
        }