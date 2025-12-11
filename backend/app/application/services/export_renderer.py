"""
Export Renderer Service
Handles template rendering and document generation (PDF/DOCX).
"""

import json
import os
from pathlib import Path
from typing import Dict, Any, Optional, List
from io import BytesIO
import zipfile

from jinja2 import Environment, FileSystemLoader, select_autoescape
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.domain.enums.export_format import ExportFormat
from app.domain.enums.template_type import TemplateType


class ExportRenderer:
    """Renders structured content into PDF/DOCX using templates."""
    
    def __init__(self):
        """Initialize the renderer with Jinja2 environment."""
        template_dir = Path(__file__).parent / "templates"
        self.env = Environment(
            loader=FileSystemLoader(str(template_dir)),
            autoescape=select_autoescape(['html', 'xml']),
            trim_blocks=True,
            lstrip_blocks=True
        )
    
    def render_pdf(
        self,
        structured_content: Dict[str, Any],
        template: TemplateType,
        options: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Render structured content to PDF using specified template.
        
        Args:
            structured_content: The structured JSON content
            template: Template type to use
            options: Template customization options (fonts, colors, etc.)
        
        Returns:
            PDF bytes
        """
        # Render HTML from template
        html_content = self._render_html(structured_content, template, options)
        
        # Convert HTML to PDF using WeasyPrint
        pdf_bytes = HTML(string=html_content).write_pdf()
        
        return pdf_bytes
    
    def render_docx(
        self,
        structured_content: Dict[str, Any],
        template: TemplateType,
        options: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Render structured content to DOCX.
        
        Args:
            structured_content: The structured JSON content
            template: Template type to use (affects styling)
            options: Template customization options
        
        Returns:
            DOCX bytes
        """
        doc = Document()
        
        # Apply template-specific styling
        style_config = self._get_docx_style_config(template, options)
        
        # Parse structured content
        data = json.loads(structured_content) if isinstance(structured_content, str) else structured_content
        header = data.get('header', {})
        sections = data.get('sections', [])
        
        # Header
        self._add_docx_header(doc, header, style_config)
        
        # Sections
        for section in sections:
            section_type = section.get('type')
            
            if section_type == 'professional_summary':
                self._add_docx_summary(doc, section, style_config)
            elif section_type == 'skills':
                self._add_docx_skills(doc, section, style_config)
            elif section_type == 'experience':
                self._add_docx_experience(doc, section, style_config)
            elif section_type == 'projects':
                self._add_docx_projects(doc, section, style_config)
            elif section_type == 'education':
                self._add_docx_education(doc, section, style_config)
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.read()
    
    def create_batch_export(
        self,
        exports: List[Dict[str, Any]]
    ) -> bytes:
        """
        Create a ZIP file containing multiple exports.
        
        Args:
            exports: List of export dictionaries with 'filename' and 'content' keys
        
        Returns:
            ZIP file bytes
        """
        buffer = BytesIO()
        
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for export in exports:
                filename = export['filename']
                content = export['content']
                zip_file.writestr(filename, content)
        
        buffer.seek(0)
        return buffer.read()
    
    def _render_html(
        self,
        structured_content: Dict[str, Any],
        template: TemplateType,
        options: Optional[Dict[str, Any]] = None
    ) -> str:
        """Render HTML from Jinja2 template."""
        # Parse structured content if string
        data = json.loads(structured_content) if isinstance(structured_content, str) else structured_content
        
        # Select template file
        template_file = f"{template.value}.html"
        jinja_template = self.env.get_template(template_file)
        
        # Merge options with defaults
        render_options = self._get_default_options(template)
        if options:
            render_options.update(options)
        
        # Render template
        html_content = jinja_template.render(
            header=data.get('header', {}),
            sections=data.get('sections', []),
            metadata=data.get('metadata', {}),
            options=render_options
        )
        
        return html_content
    
    def _get_default_options(self, template: TemplateType) -> Dict[str, Any]:
        """Get default styling options for a template."""
        defaults = {
            TemplateType.MODERN: {
                'font_family': 'Helvetica, Arial, sans-serif',
                'font_size': 11,
                'line_spacing': 1.15,
                'accent_color': '#2563EB',
            },
            TemplateType.CLASSIC: {
                'font_family': 'Georgia, Times New Roman, serif',
                'font_size': 11,
                'line_spacing': 1.2,
            },
            TemplateType.CREATIVE: {
                'font_family': 'Helvetica Neue, Arial, sans-serif',
                'font_size': 10.5,
                'line_spacing': 1.3,
                'accent_color': '#2c3e50',
                'secondary_color': '#95a5a6',
                'highlight_color': '#e74c3c',
            },
            TemplateType.ATS_OPTIMIZED: {
                'font_family': 'Arial, sans-serif',
                'font_size': 11,
                'line_spacing': 1.15,
            }
        }
        
        return defaults.get(template, {})
    
    def _get_docx_style_config(
        self,
        template: TemplateType,
        options: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Get DOCX styling configuration for a template."""
        base_config = {
            TemplateType.MODERN: {
                'heading_color': RGBColor(37, 99, 235),  # #2563EB
                'font_name': 'Calibri',
                'font_size': 11,
            },
            TemplateType.CLASSIC: {
                'heading_color': RGBColor(0, 0, 0),
                'font_name': 'Times New Roman',
                'font_size': 11,
            },
            TemplateType.CREATIVE: {
                'heading_color': RGBColor(44, 62, 80),  # #2c3e50
                'accent_color': RGBColor(231, 76, 60),  # #e74c3c
                'font_name': 'Calibri',
                'font_size': 10.5,
            },
            TemplateType.ATS_OPTIMIZED: {
                'heading_color': RGBColor(0, 0, 0),
                'font_name': 'Arial',
                'font_size': 11,
            }
        }
        
        config = base_config.get(template, base_config[TemplateType.MODERN])
        
        # Apply custom options
        if options:
            if 'font_size' in options:
                config['font_size'] = options['font_size']
            if 'font_family' in options:
                config['font_name'] = options['font_family']
        
        return config
    
    def _add_docx_header(self, doc: Document, header: Dict[str, Any], style: Dict[str, Any]):
        """Add header section to DOCX."""
        # Name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header.get('name', ''))
        run.font.size = Pt(style['font_size'] + 5)
        run.font.bold = True
        run.font.color.rgb = style['heading_color']
        
        # Title
        if header.get('title'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header['title'])
            run.font.size = Pt(style['font_size'] + 1)
        
        # Contact info
        contact_parts = []
        if header.get('location'):
            contact_parts.append(header['location'])
        if header.get('email'):
            contact_parts.append(header['email'])
        if header.get('phone'):
            contact_parts.append(header['phone'])
        
        if contact_parts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(' | '.join(contact_parts))
            run.font.size = Pt(style['font_size'] - 1)
        
        # Social links
        social_parts = []
        if header.get('linkedin'):
            social_parts.append(f"LinkedIn: {header['linkedin']}")
        if header.get('github'):
            social_parts.append(f"GitHub: {header['github']}")
        if header.get('website'):
            social_parts.append(f"Website: {header['website']}")
        
        if social_parts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(' | '.join(social_parts))
            run.font.size = Pt(style['font_size'] - 1)
        
        doc.add_paragraph()  # Spacing
    
    def _add_docx_summary(self, doc: Document, section: Dict[str, Any], style: Dict[str, Any]):
        """Add professional summary section."""
        # Section title
        p = doc.add_heading('PROFESSIONAL SUMMARY', level=1)
        p.runs[0].font.color.rgb = style['heading_color']
        
        # Content
        p = doc.add_paragraph(section.get('content', ''))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph()  # Spacing
    
    def _add_docx_skills(self, doc: Document, section: Dict[str, Any], style: Dict[str, Any]):
        """Add skills section."""
        p = doc.add_heading('SKILLS', level=1)
        p.runs[0].font.color.rgb = style['heading_color']
        
        for category in section.get('categories', []):
            # Category name
            p = doc.add_paragraph()
            run = p.add_run(f"{category['name']}: ")
            run.font.bold = True
            
            # Category items
            if category['name'] == 'Languages':
                items = [f"{lang['name']} ({lang['proficiency']})" for lang in category['items']]
                run = p.add_run(', '.join(items))
            elif category['name'] == 'Certifications':
                for cert in category['items']:
                    cert_text = f"{cert['name']} - {cert['issuer']}"
                    if cert.get('date_obtained'):
                        cert_text += f" ({cert['date_obtained']})"
                    p = doc.add_paragraph(cert_text, style='List Bullet')
            else:
                run = p.add_run(', '.join(category['items']))
        
        doc.add_paragraph()  # Spacing
    
    def _add_docx_experience(self, doc: Document, section: Dict[str, Any], style: Dict[str, Any]):
        """Add experience section."""
        p = doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
        p.runs[0].font.color.rgb = style['heading_color']
        
        for exp in section.get('entries', []):
            # Title and date
            p = doc.add_paragraph()
            run = p.add_run(exp.get('title', ''))
            run.font.bold = True
            run.font.size = Pt(style['font_size'] + 1)
            
            # Company and location
            p = doc.add_paragraph()
            company_text = exp.get('company', '')
            if exp.get('location'):
                company_text += f", {exp['location']}"
            run = p.add_run(company_text)
            run.italic = True
            
            # Date
            p = doc.add_paragraph(f"{exp.get('start_date', '')} - {exp.get('end_date', '')}")
            
            # Description
            if exp.get('description'):
                doc.add_paragraph(exp['description'])
            
            # Achievements
            for achievement in exp.get('achievements', []):
                doc.add_paragraph(achievement, style='List Bullet')
            
            doc.add_paragraph()  # Spacing between entries
    
    def _add_docx_projects(self, doc: Document, section: Dict[str, Any], style: Dict[str, Any]):
        """Add projects section."""
        p = doc.add_heading('PROJECTS', level=1)
        p.runs[0].font.color.rgb = style['heading_color']
        
        for proj in section.get('entries', []):
            # Project name and date
            p = doc.add_paragraph()
            run = p.add_run(proj.get('name', ''))
            run.font.bold = True
            run.font.size = Pt(style['font_size'] + 1)
            
            # Date
            if proj.get('start_date') or proj.get('end_date'):
                date_text = f"{proj.get('start_date', '')} - {proj.get('end_date', '')}"
                p = doc.add_paragraph(date_text)
            
            # Description
            if proj.get('description'):
                doc.add_paragraph(proj['description'])
            
            # Technologies
            if proj.get('technologies'):
                p = doc.add_paragraph()
                run = p.add_run('Technologies: ')
                run.font.bold = True
                p.add_run(', '.join(proj['technologies']))
            
            # URL
            if proj.get('url'):
                p = doc.add_paragraph()
                run = p.add_run('URL: ')
                run.font.bold = True
                p.add_run(proj['url'])
            
            doc.add_paragraph()  # Spacing
    
    def _add_docx_education(self, doc: Document, section: Dict[str, Any], style: Dict[str, Any]):
        """Add education section."""
        p = doc.add_heading('EDUCATION', level=1)
        p.runs[0].font.color.rgb = style['heading_color']
        
        for edu in section.get('entries', []):
            # Degree
            p = doc.add_paragraph()
            run = p.add_run(f"{edu.get('degree', '')} in {edu.get('field_of_study', '')}")
            run.font.bold = True
            run.font.size = Pt(style['font_size'] + 1)
            
            # Institution
            p = doc.add_paragraph()
            run = p.add_run(edu.get('institution', ''))
            run.italic = True
            
            # Date
            p = doc.add_paragraph(f"{edu.get('start_date', '')} - {edu.get('end_date', '')}")
            
            # GPA
            if edu.get('gpa'):
                p = doc.add_paragraph()
                run = p.add_run(f"GPA: {edu['gpa']}")
                run.font.bold = True
            
            # Honors
            if edu.get('honors'):
                p = doc.add_paragraph()
                run = p.add_run(', '.join(edu['honors']))
                run.italic = True
            
            doc.add_paragraph()  # Spacing
